import pandas as pd
from docx import Document
from chat_gpt_get_hyperlink import add_hyperlink

target_state="New York"

def add_paragraph(text,document,bold=False):
    print(text)
    p=document.add_paragraph()
    run=p.add_run(text)
    if(bold):
        run.bold=True

def add_link(text_before_link,link,document):
    print(f"{text_before_link}{link}")
    p=document.add_paragraph()
    p=add_hyperlink(p,link,text_before_link,link)
    

def get_useful_resources(resource_df,document):    
    filtered_df=resource_df[(resource_df["State"]==target_state)]
    add_paragraph(f"{target_state} Resources",document,bold=True)
    for index,row in filtered_df.iterrows():
        line=f"{row['Resource']}: {row['Link']}"
        add_link(row['Resource']+": ",row['Link'],document)

    add_paragraph("National Resources",document,bold=True)
    filtered_df=resource_df[(resource_df["State"]=="National")]
    for index,row in filtered_df.iterrows():
        line=f"{row['Resource']}: {row['Link']}"
        add_link(row['Resource']+": ",row['Link'],document)

    add_paragraph("",document)

def main():
    resource_df=pd.read_excel("Autism Resources by State.xlsx")
    document = Document()
    with open("intro.txt") as intro_file:
        add_paragraph(intro_file.read(),document)
    
    get_useful_resources(resource_df,document)

    with open("outro.txt") as outro_file:
        add_paragraph(outro_file.read(),document)
    document.save("Autism Resources Email.docx")

if __name__=="__main__":
    main()